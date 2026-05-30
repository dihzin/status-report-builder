from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SCREENS = "C:/Users/dirce/Pessoal/Projetos/GANTT-CHART/Gantt_Versao_Consolidada/screens"
OUTPUT = "C:/Users/dirce/Pessoal/Projetos/GANTT-CHART/Gantt_Versao_Consolidada/Gantt_Manual_v4.docx"

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

BLUE_DARK = RGBColor(0x1e, 0x40, 0xaf)
BLUE_MED  = RGBColor(0x1e, 0x3a, 0x8a)
GRAY_LIGHT = RGBColor(0x64, 0x74, 0x8b)
GRAY_PALE  = RGBColor(0x94, 0xa3, 0xb8)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BLUE_DARK
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE_MED
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    return p

def bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def img(doc, path, width_cm=15, caption=None):
    if not os.path.exists(path):
        print(f"  MISSING: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.runs[0]
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY_LIGHT
        run2.italic = True

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run("Nota:  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE_DARK
    run.italic = True
    return p

def code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    return t

# =====================================================================
# CAPA
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("Gantt Chart Interativo")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = BLUE_DARK

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_sub.add_run("Manual de Uso — versao 4")
run.font.size = Pt(16)
run.font.color.rgb = GRAY_LIGHT

doc.add_paragraph()
img(doc, f"{SCREENS}/gantt-months.png", width_cm=14,
    caption="H2R Program Roadmap — visao geral do cronograma em modo Meses")
doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run("Maio 2026  |  HTML single-file  |  Zero instalacao")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_PALE

doc.add_page_break()

# =====================================================================
# 1. VISAO GERAL
# =====================================================================
h1(doc, "1. Visao Geral da Interface")
body(doc, "O Gantt Chart Interativo e um unico arquivo HTML que roda diretamente no navegador, sem instalacao, servidor ou dependencias locais. Basta abrir gantt-v4.html no Chrome, Edge ou Firefox.")
doc.add_paragraph()
body(doc, "A interface e dividida em tres zonas:")
bullet(doc, "Toolbar (barra superior) — zoom, colapsar/expandir, busca, idioma, cores e dados.")
bullet(doc, "Painel lateral (esquerda) — tabela com ID, WBS, Tarefa, datas, Status, Progresso e Dependencias.")
bullet(doc, "Timeline Gantt (direita) — barras, dependencias, linha de hoje e marcadores.")

doc.add_paragraph()
img(doc, f"{SCREENS}/01-toolbar.png", width_cm=14,
    caption="Toolbar — controles sempre visiveis no topo da tela")
doc.add_paragraph()
note(doc, "O divisor entre painel lateral e timeline e arrastaavel. Arraste para ajustar o espaco de cada zona.")

doc.add_page_break()

# =====================================================================
# 2. PAINEL LATERAL E SWIMLANES
# =====================================================================
h1(doc, "2. Painel Lateral e Swimlanes")
body(doc, "Cada linha recebe automaticamente a cor do seu workstream (item pai de nivel 1). Isso cria swimlanes visuais que facilitam a leitura do roadmap por equipe ou frente de trabalho.")

bullet(doc, "Faixa de fundo suave na timeline (opacidade ~4-10%)")
bullet(doc, "Borda esquerda colorida de 4px no painel lateral")
bullet(doc, "Borda esquerda colorida de 4px na timeline")

doc.add_paragraph()
img(doc, f"{SCREENS}/06-status-panel.png", width_cm=11,
    caption="Painel lateral: ID, WBS, nome, datas, duracao e badges de status por workstream")

doc.add_paragraph()
h2(doc, "Coluna WBS")
body(doc, "O codigo WBS e calculado automaticamente a cada renderizacao (nao armazenado). Exemplos: 1, 1.1, 1.1.1, 1.2. Aparece no painel, no tooltip das barras e no editor — e incluido nos exports CSV/Excel.")

h2(doc, "Campo Status")
body(doc, "Controla a cor das barras de tarefa na timeline:")
make_table(doc,
    ["Valor", "Cor (tema Padrao)", "Uso sugerido"],
    [
        ("Very High", "Vermelho",       "Itens criticos, go-lives, riscos altos"),
        ("High",      "Laranja",        "Entregas importantes, desvios significativos"),
        ("Medium",    "Amarelo/Ambar",  "Execucao padrao do projeto"),
        ("Low",       "Verde",          "Tarefas de suporte, preparacao"),
        ("(vazio)",   "Cor primaria",   "Sem classificacao de prioridade"),
    ]
)

doc.add_page_break()

# =====================================================================
# 3. TIMELINE E ZOOM
# =====================================================================
h1(doc, "3. Timeline e Modos de Zoom")
body(doc, "A timeline exibe barras coloridas por status, indicadores de baseline, dependencias e marcadores de eventos. O zoom controla a granularidade temporal:")
doc.add_paragraph()
make_table(doc,
    ["Modo", "Pixels/dia", "Ideal para"],
    [
        ("Dias",    "32 px",    "Detalhe operacional, 1-4 semanas"),
        ("Semanas", "14 px",    "Tatico, 1-6 meses"),
        ("Meses",   "4 px",     "Portfolio, 6-24 meses"),
        ("Anos",    "0.6 px",   "Roadmap multi-ano"),
        ("Tudo",    "dinamico", "Fit ao container — visao panoramica completa"),
    ]
)
doc.add_paragraph()
img(doc, f"{SCREENS}/03-timeline-top.png", width_cm=14,
    caption="Timeline — workstreams com cores distintas, barras e baseline visiveis")

doc.add_page_break()

# =====================================================================
# 4. BASELINE TRACKING
# =====================================================================
h1(doc, "4. Baseline Tracking — Rastreamento de Desvio")
body(doc, "Os campos baseline_start e baseline_end permitem comparar o plano original com a execucao atual. Quando preenchidos, o grafico exibe:")
bullet(doc, "Barra cinza translucida — representa o periodo planejado originalmente")
bullet(doc, "Barra colorida atual — sobreposicao com a execucao real")
bullet(doc, "Badge ambar +Nw — indica atraso em semanas quando fim atual > baseline_end")

doc.add_paragraph()
img(doc, f"{SCREENS}/04-baseline-badges.png", width_cm=14,
    caption="Baseline tracking: barra cinza (plano) + badges +3w e +17w indicando semanas de atraso")

doc.add_paragraph()
h2(doc, "Como adicionar baseline")
body(doc, "No editor Gerenciar Dados, preencha baseline_start e baseline_end com as datas do plano original. Via CSV/Excel, adicione as colunas com esses nomes. Exemplo JSON:")
code_line(doc, '{ "baseline_start": "2026-06-29", "baseline_end": "2026-07-17" }')
doc.add_paragraph()
note(doc, "O badge +Nw aparece apenas quando o desvio e positivo (atraso). Projetos adiantados nao exibem badge.")

doc.add_page_break()

# =====================================================================
# 5. MARCADORES DE EVENTOS
# =====================================================================
h1(doc, "5. Marcadores de Eventos — Campo marker")
body(doc, "O campo marker adiciona uma linha vertical destacada na timeline sobre a data de inicio da tarefa, com o nome rotacionado verticalmente ao longo da linha. Ideal para go-lives e gates de decisao.")

doc.add_paragraph()
make_table(doc,
    ["Valor", "Visual na Timeline", "Uso"],
    [
        ('"go-live"',   "Linha azul solida + triangulo azul A",      "Go Live, ativacao de sistema, entrada em producao"),
        ('"decision"',  "Linha vermelha tracejada + losango vermelho", "Go/No Go, gate de aprovacao, checkpoint critico"),
        ('"" (vazio)',  "Sem linha vertical",                         "Comportamento padrao — marco normal"),
    ]
)
doc.add_paragraph()
img(doc, f"{SCREENS}/05-marker-lines.png", width_cm=14,
    caption="Marcadores na timeline: linhas azuis (go-live) e linha vermelha tracejada (Go/No Go) com rotulos verticais")

doc.add_paragraph()
h2(doc, "Como adicionar um marker")
body(doc, "1. Abrir Gerenciar Dados na toolbar.")
body(doc, "2. Localizar a linha da tarefa desejada.")
body(doc, "3. No campo marker, digitar go-live ou decision.")
body(doc, "4. Clicar Aplicar.")
doc.add_paragraph()
code_line(doc, '{ "marker": "go-live", "milestone": true, "start": "2026-08-31", "end": "2026-08-31" }')
doc.add_paragraph()
note(doc, "O texto exibido na linha e o nome da tarefa, truncado em 22 caracteres. Use nomes curtos e diretos para tarefas com marker.")

doc.add_page_break()

# =====================================================================
# 6. SISTEMA DE CORES
# =====================================================================
h1(doc, "6. Sistema de Cores — Dois Sistemas Independentes")
body(doc, "O botao de paleta (icone de pincel) na toolbar controla dois sistemas completamente independentes que podem ser combinados livremente:")

doc.add_paragraph()
h2(doc, "6.1 Color Theme — Cor das Barras por Status")
body(doc, "Controla a cor das barras de tarefa de acordo com o campo status. Temas disponiveis: Padrao, Vibrante, Pastel, Oceano, Floresta, Mono Cinza e Corporativo.")
body(doc, "E possivel personalizar as 6 cores individualmente (Muito Alta, Alta, Media, Baixa, Fase, Primaria) e salvar como tema proprio clicando em Salvar como...")

doc.add_paragraph()
h2(doc, "6.2 Paleta Workstreams — Faixas de Fundo e Bordas")
body(doc, "Controla as cores das swimlanes. Tres presets disponiveis:")
bullet(doc, "Colorido — azul, violeta, ciano, verde, ambar, vermelho, rosa, laranja (ideal para apresentacoes)")
bullet(doc, "Pastel — versoes suaves das mesmas cores (ambientes sobrios, fundo claro)")
bullet(doc, "Neutro — tons de cinza (documentos formais, impressao)")

doc.add_paragraph()
img(doc, f"{SCREENS}/07-theme-popover.png", width_cm=8,
    caption="Painel de Cores: Tema de Cores (presets de status) acima, Paleta Workstreams abaixo")

doc.add_paragraph()
note(doc, "Os dois sistemas sao independentes. Qualquer Color Theme pode ser combinado com qualquer Paleta Workstreams. As escolhas sao salvas automaticamente.")

doc.add_page_break()

# =====================================================================
# 7. GERENCIAR DADOS
# =====================================================================
h1(doc, "7. Gerenciar Dados")
body(doc, "O botao verde Gerenciar Dados na toolbar abre o editor central do cronograma. A partir dele, e possivel editar, importar, exportar e persistir todos os dados do projeto.")

doc.add_paragraph()
img(doc, f"{SCREENS}/08-manage-data.png", width_cm=14,
    caption="Editor Gerenciar Dados — tabela inline com drag & drop, indent/outdent e historico de desfazer")

doc.add_paragraph()
h2(doc, "Editor inline")
bullet(doc, "Drag & drop de linhas para reordenar")
bullet(doc, "Indent (tabela) / Outdent: altera o parent_id — muda a hierarquia da tarefa")
bullet(doc, "Desfazer: historico de ate 50 estados")
bullet(doc, "Duplicar / Excluir linha individualmente")
bullet(doc, "Aplicar: aplica ao grafico e salva no localStorage")

h2(doc, "Import CSV/Excel")
body(doc, "Clique em Importar CSV/Excel, selecione o arquivo (.csv ou .xlsx), revise o preview e escolha Substituir ou Adicionar. Os cabecalhos sao reconhecidos automaticamente em PT-BR, ES e EN.")

h2(doc, "Export")
make_table(doc,
    ["Botao", "Saida gerada"],
    [
        ("CSV",          "gantt-data.csv — com coluna WBS incluida"),
        ("Excel",        "gantt-data.xlsx — com coluna WBS incluida"),
        ("PNG",          "gantt-chart.png — imagem completa (painel + timeline)"),
        ("Copiar JSON",  "JSON dos dados copiado para o clipboard"),
        ("Salvar",       "Forca gravacao imediata no localStorage"),
    ]
)

doc.add_page_break()

# =====================================================================
# 8. FILTROS E BUSCA
# =====================================================================
h1(doc, "8. Filtros e Busca")

h2(doc, "View Manager")
body(doc, "Clique em Visualizacao na toolbar para abrir o View Manager:")
bullet(doc, "Mostrar/ocultar e redimensionar colunas do painel lateral")
bullet(doc, "Filtrar: ocultar marcos, ocultar 100% concluidos, ocultar fins de semana, ocultar % nas barras")
bullet(doc, "Filtros avancados: por status, faixa de % de conclusao, responsavel e janela de datas")
bullet(doc, "Salvar visualizacoes nomeadas para alternar rapidamente entre vistas")

h2(doc, "Campo de Busca")
body(doc, "O campo Buscar... na toolbar destaca linhas com o texto buscado: fundo amarelo pulsante na linha do painel e borda laranja na barra da timeline. Limpa ao apagar o texto.")

doc.add_page_break()

# =====================================================================
# 9. PERSISTENCIA E IDIOMAS
# =====================================================================
h1(doc, "9. Persistencia e Idiomas")

h2(doc, "Dados e configuracoes salvos automaticamente")
body(doc, "O grafico usa o localStorage do navegador (chave gantt-data-v3) para persistir:")
bullet(doc, "Dados das tarefas (apos clicar Aplicar ou Salvar)")
bullet(doc, "Configuracoes de colunas — visibilidade e largura")
bullet(doc, "Tema de cores e paleta de workstreams selecionados")
bullet(doc, "Idioma selecionado")
bullet(doc, "Visualizacoes salvas")
doc.add_paragraph()
note(doc, "Atencao: limpar dados do navegador apaga o projeto. Faca backup regular usando Export CSV ou Excel.")

h2(doc, "Suporte a Idiomas")
body(doc, "O seletor PT-BR / ES / EN no canto direito da toolbar traduz toda a interface — botoes, modais, tooltips, cabecalhos da timeline e mensagens de erro. A preferencia e salva no localStorage.")

doc.add_page_break()

# =====================================================================
# 10. MODELO DE DADOS
# =====================================================================
h1(doc, "10. Modelo de Dados Completo")
body(doc, "Cada tarefa e um objeto JSON com os seguintes campos. Os marcados com * sao obrigatorios:")
doc.add_paragraph()
make_table(doc,
    ["Campo", "Tipo", "Descricao"],
    [
        ("id *",             "string",  "Identificador unico (ex: \"1\", \"DEV\")"),
        ("parent_id *",      "string",  "ID do pai na hierarquia; \"\" para raiz"),
        ("task *",           "string",  "Nome da tarefa ou fase"),
        ("start *",          "string",  "Data de inicio YYYY-MM-DD"),
        ("end *",            "string",  "Data de termino YYYY-MM-DD"),
        ("milestone",        "boolean", "true = marco (diamante/triangulo)"),
        ("completion",       "number",  "Percentual concluido, 0-100"),
        ("dependencies",     "string",  "IDs predecessores separados por virgula"),
        ("assignee",         "string",  "Nome do responsavel"),
        ("status",           "string",  "\"Very High\", \"High\", \"Medium\", \"Low\" ou \"\""),
        ("marker",           "string",  "\"go-live\", \"decision\" ou \"\" (padrao)"),
        ("baseline_start",   "string",  "Inicio do plano original YYYY-MM-DD"),
        ("baseline_end",     "string",  "Fim do plano original YYYY-MM-DD"),
    ]
)

doc.add_paragraph()
note(doc, "Os campos baseline_start e baseline_end, quando preenchidos, ativam o baseline tracking (barra cinza + badge +Nw) automaticamente sem nenhuma configuracao adicional.")

# =====================================================================
# SAVE
# =====================================================================
doc.save(OUTPUT)
print(f"Word document saved: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT):,} bytes")
